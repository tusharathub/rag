import os
import hashlib
import mimetypes
import json
from typing import BinaryIO, Dict, Any, Union, List, Tuple
from pydantic import BaseModel
from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook
from pptx import Presentation
from bs4 import BeautifulSoup

# Define supported extensions and MIME types
SUPPORTED_FORMATS: Dict[str, List[str]] = {
    "pdf": ["application/pdf"],
    "docx": ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
    "xlsx": ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"],
    "xls": ["application/vnd.ms-excel"],
    "pptx": ["application/vnd.openxmlformats-officedocument.presentationml.presentation"],
    "ppt": ["application/vnd.ms-powerpoint"],
    "txt": ["text/plain"],
    "md": ["text/markdown", "text/x-markdown", "text/plain"],
    "csv": ["text/csv", "application/csv", "text/plain"],
    "html": ["text/html"],
    "htm": ["text/html"],
    "json": ["application/json", "text/plain"]
}

# Map extensions to their magic bytes
MAGIC_NUMBERS: Dict[str, bytes] = {
    "pdf": b"%PDF-",
    "docx": b"PK\x03\x04",
    "xlsx": b"PK\x03\x04",
    "pptx": b"PK\x03\x04",
    "xls": b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1",
    "ppt": b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1",
}


class FileValidationError(Exception):
    """Custom exception raised when file validation fails."""
    pass


def compute_sha256(file_path: str) -> str:
    """Compute the SHA-256 hash of a file on disk in a streaming manner."""
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        # Read in 64KB chunks to optimize memory usage
        for byte_block in iter(lambda: f.read(65536), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()


def validate_file_type(file_path: str, filename: str, content_type: str) -> Tuple[str, str]:
    """Validates the file extension, MIME type, and magic bytes.
    
    Args:
        file_path: Path to the temporary file on disk.
        filename: Original name of the uploaded file.
        content_type: MIME type reported by the upload request.
        
    Returns:
        Tuple[str, str]: Normalized extension and normalized MIME type.
        
    Raises:
        FileValidationError: If verification fails.
    """
    # 1. Extract and check extension
    _, ext = os.path.splitext(filename.lower())
    ext = ext.lstrip(".")
    if not ext:
        raise FileValidationError("File has no extension.")
        
    if ext not in SUPPORTED_FORMATS:
        raise FileValidationError(f"Unsupported file extension: .{ext}")

    # 2. Check reported MIME type compatibility
    allowed_mimes = SUPPORTED_FORMATS[ext]
    # Allow empty/generic octet-stream reported by some clients, but prioritize validation
    if content_type and content_type != "application/octet-stream" and content_type not in allowed_mimes:
        # Sometimes txt/md/csv are reported as text/plain or vice-versa, which we support
        if not (ext in ["txt", "md", "csv", "json"] and content_type.startswith("text/")):
            raise FileValidationError(
                f"MIME type mismatch. Reported: {content_type}, Expected one of: {allowed_mimes}"
            )

    # 3. Check magic bytes (for binary files)
    try:
        with open(file_path, "rb") as f:
            header_bytes = f.read(8)
    except Exception as e:
        raise FileValidationError(f"Could not read file header for validation: {str(e)}")

    if ext in MAGIC_NUMBERS:
        expected_prefix = MAGIC_NUMBERS[ext]
        if not header_bytes.startswith(expected_prefix):
            # Check if this is an Office XML file which might share PK\x03\x04
            if expected_prefix == b"PK\x03\x04" and header_bytes.startswith(b"PK\x03\x04"):
                pass
            else:
                raise FileValidationError(
                    f"File signature (magic bytes) does not match the extension .{ext}."
                )
    else:
        # For text files, check if it can be decoded as UTF-8
        try:
            with open(file_path, "r", encoding="utf-8", errors="strict") as f:
                # Read first 8KB to verify text decoding
                f.read(8192)
        except UnicodeDecodeError:
            raise FileValidationError(
                f"File contents for .{ext} are not valid UTF-8 text."
            )

    # Normalize response MIME type
    final_mime = allowed_mimes[0] if allowed_mimes else "application/octet-stream"
    return ext, final_mime


def extract_metadata(file_path: str, file_type: str) -> Dict[str, str]:
    """Extracts metadata from the file based on its type.
    
    Args:
        file_path: Path to the file on disk.
        file_type: Extension name (e.g. 'pdf', 'docx').
        
    Returns:
        Dict[str, str]: Key-value string pairs representing metadata.
    """
    metadata: Dict[str, str] = {}
    
    try:
        if file_type == "pdf":
            reader = PdfReader(file_path)
            metadata["page_count"] = str(len(reader.pages))
            if reader.metadata:
                for k, v in reader.metadata.items():
                    clean_key = k.lstrip("/")
                    metadata[clean_key] = str(v)
                    
        elif file_type == "docx":
            doc = DocxDocument(file_path)
            metadata["paragraph_count"] = str(len(doc.paragraphs))
            metadata["table_count"] = str(len(doc.tables))
            
            props = doc.core_properties
            for prop in ["author", "category", "created", "last_modified_by", "modified", "title", "subject"]:
                try:
                    val = getattr(props, prop)
                    if val:
                        metadata[prop] = str(val)
                except Exception:
                    pass
                    
        elif file_type == "xlsx" or file_type == "xls":
            # For xls, we just treat sheet count using openpyxl for xlsx, but catch exceptions for old xls
            if file_type == "xlsx":
                wb = load_workbook(file_path, read_only=True)
                metadata["sheet_count"] = str(len(wb.sheetnames))
                metadata["sheet_names"] = ", ".join(wb.sheetnames)
            else:
                metadata["sheet_count"] = "unknown (legacy XLS)"
                
        elif file_type == "pptx" or file_type == "ppt":
            if file_type == "pptx":
                prs = Presentation(file_path)
                metadata["slide_count"] = str(len(prs.slides))
                try:
                    props = prs.core_properties
                    for prop in ["author", "created", "last_modified_by", "modified", "title"]:
                        val = getattr(props, prop)
                        if val:
                            metadata[prop] = str(val)
                except Exception:
                    pass
            else:
                metadata["slide_count"] = "unknown (legacy PPT)"
                
        elif file_type == "html" or file_type == "htm":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
            metadata["title"] = soup.title.string if soup.title else ""
            desc_tag = soup.find("meta", attrs={"name": "description"})
            if desc_tag and desc_tag.get("content"):
                metadata["description"] = desc_tag.get("content")
            metadata["h1_count"] = str(len(soup.find_all("h1")))
            metadata["h2_count"] = str(len(soup.find_all("h2")))
            
        elif file_type == "json":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                data = json.load(f)
            metadata["type"] = "array" if isinstance(data, list) else "object"
            metadata["item_count"] = str(len(data)) if isinstance(data, (list, dict)) else "0"
            
        elif file_type in ["txt", "md", "csv"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            metadata["line_count"] = str(len(content.splitlines()))
            metadata["word_count"] = str(len(content.split()))
            metadata["character_count"] = str(len(content))
            
    except Exception as e:
        # Don't fail the whole upload if metadata extraction fails, just record the error in metadata
        metadata["extraction_error"] = str(e)
        
    return metadata
